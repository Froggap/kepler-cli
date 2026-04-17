from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class WordService:
    PRIMARY_COLOR   = "000000"
    SECONDARY_COLOR = "2D2D2D"
    WHITE_COLOR     = "FFFFFF"

    def __init__(self, data: dict):
        self.data = data
        self.doc  = Document()
        self._setup_margins()
        self._setup_base_style()          

    def _setup_margins(self):
        for section in self.doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(3)
            section.right_margin  = Cm(3)

    def _setup_base_style(self):         
        style = self.doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor(0, 0, 0)

    def _set_cell_bg(self, cell, hex_color: str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _add_heading(self, text: str, level: int = 1, color: str = None):
        color = color or self.PRIMARY_COLOR        
        p = self.doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(color)
        return p

    def _add_bullet(self, text: str):
        p   = self.doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.size = Pt(11)
        return p

    def _add_header_row(self, table, labels: list[str], color: str = None):
        color = color or self.PRIMARY_COLOR       
        row   = table.rows[0]
        for i, label in enumerate(labels):
            cell = row.cells[i]
            cell.text = label
            self._set_cell_bg(cell, color)
            run = cell.paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor.from_string(self.WHITE_COLOR)  
            run.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _spacer(self):
        self.doc.add_paragraph()

    def _build_header(self):
        enc = self.data["encabezado"]

        title = self.doc.add_paragraph()
        run = title.add_run(enc.get("empresa", "").upper())
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor.from_string(self.PRIMARY_COLOR)     

        sub = self.doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run("Reporte Mensual de Actividades").font.size = Pt(13)
        

        self._spacer()

        col = self.doc.add_paragraph()
        run = col.add_run("Colaborador: ")
        run.bold = True
        run.font.size = Pt(11)
        col.add_run(enc.get('colaborador', '')).font.size = Pt(11)

        proy = self.doc.add_paragraph()
        run = proy.add_run("Proyecto: ")
        run.bold = True
        run.font.size = Pt(11)
        proy.add_run(enc.get('proyecto', '')).font.size = Pt(11)

        per = self.doc.add_paragraph()
        run = per.add_run("Período: ")
        run.bold = True
        run.font.size = Pt(11)
        per.add_run(enc.get('periodo', '')).font.size = Pt(11)

        self._spacer()

    def _build_summary(self):
        self._add_heading("Resumen Ejecutivo", level=1)
        for parrafo in self.data.get("resumen_ejecutivo", []):
            p = self.doc.add_paragraph(parrafo)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
        self._spacer()

    def _build_activities(self):
        self._add_heading("Actividades Realizadas", level=1)
        for categoria in self.data.get("actividades", []):
            self._add_heading(categoria["categoria"], level=2, color=self.SECONDARY_COLOR)  
            for item in categoria.get("items", []):
                self._add_bullet(item)
        self._spacer()

    def _build_results(self):
        self._add_heading("Resultados y Valor Generado", level=1)
        for resultado in self.data.get("resultados_valor", []):
            self._add_bullet(resultado)
        self._spacer()


    def build(self) -> "WordService":
        print("🛠  Construyendo reporte en Word...")
        self._build_header()
        self._build_summary()
        self._build_activities()
        self._build_results()
        return self

    def save(self, output_path: str) -> None:
        while True:
            try:
                self.doc.save(output_path)
                print(f"✅ Reporte guardado en: {output_path}")
                break;
            except PermissionError:
                input(f"⚠️  No se pudo guardar el archivo. Cierra '{output_path}' y presiona Enter para reintentar...")